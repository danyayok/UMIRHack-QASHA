import os
import json
from datetime import datetime
from typing import List, Dict, Any
import logging
from pathlib import Path

logger = logging.getLogger("qa_automata")


class TestCaseExporter:
    def __init__(self):
        self.export_dir = "./storage/exports"
        os.makedirs(self.export_dir, exist_ok=True)

    async def export_test_cases(self, test_cases: List, format: str) -> Dict[str, Any]:
        """Экспортирует тест-кейсы в указанном формате"""
        try:
            if format == "excel":
                return await self._export_to_excel(test_cases)
            elif format == "word":
                return await self._export_to_word(test_cases)
            elif format == "txt":
                return await self._export_to_txt(test_cases)
            else:
                raise ValueError(f"Unsupported format: {format}")

        except Exception as e:
            logger.error(f"Export error: {e}")
            raise

    async def _export_to_excel(self, test_cases: List) -> Dict[str, Any]:
        """Экспорт в Excel формате"""
        try:
            import pandas as pd

            # Подготавливаем данные для Excel
            data = []
            for tc in test_cases:
                # Преобразуем шаги в строку
                steps_text = ""
                if tc.steps:
                    for step in tc.steps:
                        steps_text += f"{step.get('step_number', 1)}. {step.get('action', '')}\n"

                data.append({
                    "ID Тест-кейса": tc.test_case_id,
                    "Название": tc.name,
                    "Описание": tc.description or "",
                    "Тип": tc.test_type,
                    "Приоритет": tc.priority,
                    "Шаги": steps_text,
                    "Предусловия": tc.preconditions or "",
                    "Постусловия": tc.postconditions or "",
                    "Статус": tc.status
                })

            df = pd.DataFrame(data)

            # Сохраняем в файл
            filename = f"test_cases_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            filepath = os.path.join(self.export_dir, filename)

            df.to_excel(filepath, index=False, engine='openpyxl')

            return {
                "filename": filename,
                "filepath": filepath,
                "format": "excel",
                "test_cases_count": len(test_cases),
                "download_url": f"/api/v1/download/{filename}"
            }

        except ImportError:
            logger.error("pandas or openpyxl not installed")
            raise Exception("Excel export requires pandas and openpyxl")

    async def _export_to_word(self, test_cases: List) -> Dict[str, Any]:
        """Экспорт в Word формате"""
        try:
            from docx import Document
            from docx.shared import Inches

            doc = Document()

            # Заголовок
            doc.add_heading('Тест-кейсы', 0)
            doc.add_paragraph(f'Дата экспорта: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
            doc.add_paragraph(f'Количество тест-кейсов: {len(test_cases)}')
            doc.add_paragraph()

            for tc in test_cases:
                # Заголовок тест-кейса
                doc.add_heading(f'{tc.test_case_id}: {tc.name}', level=2)

                # Основная информация
                doc.add_paragraph(f'Тип: {tc.test_type}')
                doc.add_paragraph(f'Приоритет: {tc.priority}')
                doc.add_paragraph(f'Статус: {tc.status}')

                if tc.description:
                    doc.add_paragraph('Описание:')
                    doc.add_paragraph(tc.description)

                if tc.preconditions:
                    doc.add_paragraph('Предусловия:')
                    doc.add_paragraph(tc.preconditions)

                # Шаги
                if tc.steps:
                    doc.add_paragraph('Шаги выполнения:')
                    for step in tc.steps:
                        p = doc.add_paragraph()
                        p.add_run(f"{step.get('step_number', 1)}. ").bold = True
                        p.add_run(f"{step.get('action', '')}")
                        if step.get('expected_result'):
                            p.add_run(f" → {step.get('expected_result')}")

                if tc.postconditions:
                    doc.add_paragraph('Постусловия:')
                    doc.add_paragraph(tc.postconditions)

                doc.add_paragraph()  # Пустая строка между тест-кейсами

            # Сохраняем файл
            filename = f"test_cases_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join(self.export_dir, filename)
            doc.save(filepath)

            return {
                "filename": filename,
                "filepath": filepath,
                "format": "word",
                "test_cases_count": len(test_cases),
                "download_url": f"/api/v1/download/{filename}"
            }

        except ImportError:
            logger.error("python-docx not installed")
            raise Exception("Word export requires python-docx")

    async def _export_to_txt(self, test_cases: List) -> Dict[str, Any]:
        """Экспорт в TXT формате"""
        try:
            filename = f"test_cases_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
            filepath = os.path.join(self.export_dir, filename)

            with open(filepath, 'w', encoding='utf-8') as f:
                f.write("ТЕСТ-КЕЙСЫ\n")
                f.write("=" * 50 + "\n")
                f.write(f"Дата экспорта: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n")
                f.write(f"Количество тест-кейсов: {len(test_cases)}\n\n")

                for tc in test_cases:
                    f.write(f"{tc.test_case_id}: {tc.name}\n")
                    f.write("-" * 50 + "\n")
                    f.write(f"Тип: {tc.test_type}\n")
                    f.write(f"Приоритет: {tc.priority}\n")
                    f.write(f"Статус: {tc.status}\n")

                    if tc.description:
                        f.write(f"Описание: {tc.description}\n")

                    if tc.preconditions:
                        f.write(f"Предусловия: {tc.preconditions}\n")

                    if tc.steps:
                        f.write("Шаги выполнения:\n")
                        for step in tc.steps:
                            f.write(f"  {step.get('step_number', 1)}. {step.get('action', '')}")
                            if step.get('expected_result'):
                                f.write(f" → {step.get('expected_result')}")
                            f.write("\n")

                    if tc.postconditions:
                        f.write(f"Постусловия: {tc.postconditions}\n")

                    f.write("\n")  # Пустая строка между тест-кейсами

            return {
                "filename": filename,
                "filepath": filepath,
                "format": "txt",
                "test_cases_count": len(test_cases),
                "download_url": f"/api/v1/download/{filename}"
            }

        except Exception as e:
            logger.error(f"TXT export error: {e}")
            raise